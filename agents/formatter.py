# agents/formatter.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def build_docx(resume_text: str, output_path: str = "optimized_resume.docx") -> str:
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = resume_text.strip().split('\n')
    bullet_pattern = re.compile(r'^[-*•]\s+(.*)')
    section_title_pattern = re.compile(r'^\s*\*\*(.+?)\*\*\s*$')

    for line in lines:
        line = line.strip()

        if not line:
            document.add_paragraph("")
            continue

        # Section titles
        section_match = section_title_pattern.match(line)
        if section_match:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(section_match.group(1))
            run.bold = True
            run.font.size = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Bullet points
        bullet_match = bullet_pattern.match(line)
        if bullet_match:
            paragraph = document.add_paragraph(style='List Bullet')
            run = paragraph.add_run(bullet_match.group(1))
            run.font.size = Pt(11)
            continue

        # Regular paragraph
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)

    document.save(output_path)
    return output_path
